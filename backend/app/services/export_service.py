import io
import tempfile
from typing import BinaryIO
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
import re
from app.db.models import GeneratedDocument

class DocumentExportService:
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        # Create custom styles
        self.title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Title'],
            fontSize=18,
            spaceAfter=20,
            alignment=TA_CENTER
        )
        self.heading_style = ParagraphStyle(
            'CustomHeading',
            parent=self.styles['Heading1'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=12
        )
        self.normal_style = ParagraphStyle(
            'CustomNormal',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            alignment=TA_LEFT
        )
    
    def markdown_to_plain_text(self, markdown_text: str) -> str:
        """Convert markdown to plain text for PDF generation"""
        # Convert markdown to HTML first
        html = markdown.markdown(markdown_text)
        # Remove HTML tags
        clean = re.compile('<.*?>')
        return re.sub(clean, '', html)
    
    def export_to_pdf(self, generated_doc: GeneratedDocument) -> io.BytesIO:
        """Export a generated document to PDF format"""
        buffer = io.BytesIO()
        
        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Build story (content)
        story = []
        
        # Title
        title = f"{generated_doc.title}"
        story.append(Paragraph(title, self.title_style))
        story.append(Spacer(1, 20))
        
        # Document type and metadata
        meta_info = f"Document Type: {generated_doc.document_type.upper()}<br/>"
        meta_info += f"Created: {generated_doc.created_at.strftime('%B %d, %Y')}<br/>"
        meta_info += f"Version: {generated_doc.version}"
        story.append(Paragraph(meta_info, self.normal_style))
        story.append(Spacer(1, 20))
        
        # Content
        # Convert markdown content to plain text and split into paragraphs
        plain_content = self.markdown_to_plain_text(generated_doc.content)
        paragraphs = plain_content.split('\n\n')
        
        for para in paragraphs:
            if para.strip():
                # Check if it's a heading (starts with #)
                if para.strip().startswith('#'):
                    # Remove # symbols and treat as heading
                    heading_text = para.strip().lstrip('#').strip()
                    story.append(Paragraph(heading_text, self.heading_style))
                else:
                    story.append(Paragraph(para.strip(), self.normal_style))
                story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def export_to_word(self, generated_doc: GeneratedDocument) -> io.BytesIO:
        """Export a generated document to Word format"""
        doc = Document()
        
        # Title
        title_para = doc.add_heading(generated_doc.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Document Type: {generated_doc.document_type.upper()}")
        doc.add_paragraph(f"Created: {generated_doc.created_at.strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Version: {generated_doc.version}")
        doc.add_paragraph("")  # Empty line
        
        # Content - convert markdown to Word
        content_lines = generated_doc.content.split('\n')
        current_paragraph = ""
        
        for line in content_lines:
            line = line.strip()
            
            if not line:  # Empty line
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)
                    current_paragraph = ""
                doc.add_paragraph("")  # Add empty paragraph
                continue
            
            # Handle headers
            if line.startswith('#'):
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)
                    current_paragraph = ""
                
                # Count # symbols to determine header level
                level = 0
                while level < len(line) and line[level] == '#':
                    level += 1
                
                header_text = line[level:].strip()
                doc.add_heading(header_text, level=min(level, 3))
                continue
            
            # Handle list items
            if line.startswith('-') or line.startswith('*') or line.startswith('+'):
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)
                    current_paragraph = ""
                
                # Remove list marker and add as paragraph with bullet
                list_text = line[1:].strip()
                doc.add_paragraph(list_text, style='List Bullet')
                continue
            
            # Regular text - accumulate into paragraph
            if current_paragraph:
                current_paragraph += " " + line
            else:
                current_paragraph = line
        
        # Add final paragraph if exists
        if current_paragraph:
            doc.add_paragraph(current_paragraph)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def get_filename(self, generated_doc: GeneratedDocument, format_type: str) -> str:
        """Generate appropriate filename for export"""
        # Sanitize title for filename
        safe_title = re.sub(r'[^\w\s-]', '', generated_doc.title)
        safe_title = re.sub(r'[-\s]+', '-', safe_title)
        
        extension = 'pdf' if format_type.lower() == 'pdf' else 'docx'
        return f"{safe_title}-{generated_doc.document_type}-v{generated_doc.version}.{extension}"